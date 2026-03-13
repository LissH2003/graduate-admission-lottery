#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成北科大研究生复试抽签系统用户操作手册（Word版）
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._element.tcPr
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_elm = tc.find(qn(f'w:{edge}'))
            if edge_elm is None:
                edge_elm = docx.oxml.OxmlElement(f'w:{edge}')
                tc.append(edge_elm)
            edge_elm.set(qn('w:val'), kwargs[edge])

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_screenshot_placeholder(doc, title, description=""):
    """添加截图占位符"""
    # 添加灰色背景的提示框
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = paragraph.add_run(f"【截图位置：{title}】")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.bold = True
    
    if description:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(description)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(150, 150, 150)
    
    # 添加空行作为截图区域
    doc.add_paragraph()
    doc.add_paragraph()

def main():
    doc = Document()
    
    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    
    # ===== 封面 =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('北京科技大学\n研究生复试抽签系统\n用户操作手册')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    v_run = version.add_run('版本：V2.1.0\n日期：2026年3月')
    v_run.font.size = Pt(14)
    
    doc.add_page_break()
    
    # ===== 目录 =====
    add_heading_custom(doc, '目录', level=1)
    toc_items = [
        '一、系统访问与登录',
        '    1.1 访问系统',
        '    1.2 统一身份认证登录',
        '    1.3 系统账号登录',
        '二、管理员操作指南',
        '    2.1 管理员首页',
        '    2.2 批次管理',
        '    2.3 考场管理',
        '    2.4 考生导入',
        '    2.5 抽签执行',
        '三、志愿者操作指南',
        '    3.1 选择考场',
        '    3.2 查看分组信息',
        '四、常见问题',
        '    4.1 登录问题',
        '    4.2 操作问题'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 第一章：系统访问与登录 =====
    add_heading_custom(doc, '一、系统访问与登录', level=1)
    
    # 1.1
    add_heading_custom(doc, '1.1 访问系统', level=2)
    doc.add_paragraph('推荐使用 Chrome 或 Edge 浏览器访问系统。')
    doc.add_paragraph('系统地址：http://115.25.59.77')
    doc.add_paragraph()
    add_screenshot_placeholder(doc, '图1-1：系统登录首页', '显示两种登录方式：统一身份认证登录、系统账号登录')
    
    # 1.2
    add_heading_custom(doc, '1.2 统一身份认证登录', level=2)
    steps = [
        '点击【统一身份认证登录】按钮',
        '系统自动跳转至学校竹云平台',
        '在学校认证页面输入您的工号/学号和密码',
        '认证成功后自动返回本系统',
        '根据您的权限进入管理员首页或志愿者页面'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    add_screenshot_placeholder(doc, '图1-2：统一身份认证页面', '竹云学校认证页面，包含工号/学号、密码输入框')
    
    doc.add_paragraph('【注意事项】')
    doc.add_paragraph('• 首次使用统一身份认证，请确保已在信息中心开通权限', style='List Bullet')
    doc.add_paragraph('• 如忘记密码，请联系学校信息中心重置', style='List Bullet')
    
    # 1.3
    add_heading_custom(doc, '1.3 系统账号登录（备用）', level=2)
    doc.add_paragraph('当统一身份认证不可用时，可使用系统账号登录。')
    doc.add_paragraph()
    
    steps2 = [
        '点击【系统账号登录】按钮',
        '选择您的角色：管理员 或 志愿者',
        '输入用户名和密码',
        '点击【登录系统】'
    ]
    for i, step in enumerate(steps2, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    add_screenshot_placeholder(doc, '图1-3：系统账号登录界面', '显示角色选择、用户名密码输入框、登录按钮')
    
    doc.add_paragraph()
    doc.add_paragraph('【测试账号】')
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '角色'
    hdr_cells[1].text = '用户名'
    hdr_cells[2].text = '密码'
    
    row1 = table.rows[1].cells
    row1[0].text = '管理员'
    row1[1].text = 'admin'
    row1[2].text = 'admin123'
    
    row2 = table.rows[2].cells
    row2[0].text = '志愿者'
    row2[1].text = 'volunteer1'
    row2[2].text = 'volunteer123'
    
    doc.add_page_break()
    
    # ===== 第二章：管理员操作指南 =====
    add_heading_custom(doc, '二、管理员操作指南', level=1)
    
    # 2.1
    add_heading_custom(doc, '2.1 管理员首页', level=2)
    doc.add_paragraph('登录成功后进入管理员首页，可查看系统概览和快捷入口。')
    add_screenshot_placeholder(doc, '图2-1：管理员首页', '显示功能导航菜单、统计信息、快捷操作按钮')
    
    # 2.2
    add_heading_custom(doc, '2.2 批次管理', level=2)
    doc.add_paragraph('批次是复试抽签的基本单位，包含时间、地点等基本信息。')
    doc.add_paragraph('【操作步骤】')
    steps3 = [
        '点击左侧菜单【批次管理】',
        '点击【新建批次】按钮',
        '填写批次名称、时间、备注信息',
        '点击【保存】完成创建',
        '可在列表中启用/停用批次'
    ]
    for i, step in enumerate(steps3, 1):
        doc.add_paragraph(f'{i}. {step}')
    add_screenshot_placeholder(doc, '图2-2：批次管理页面', '显示批次列表、新建按钮、状态开关')
    
    # 2.3
    add_heading_custom(doc, '2.3 考场管理', level=2)
    doc.add_paragraph('配置复试考场信息，设置考场容量。')
    doc.add_paragraph('【操作步骤】')
    steps4 = [
        '点击左侧菜单【考场管理】',
        '点击【添加考场】',
        '填写考场名称、地址、容量',
        '关联到对应批次',
        '点击【保存】'
    ]
    for i, step in enumerate(steps4, 1):
        doc.add_paragraph(f'{i}. {step}')
    add_screenshot_placeholder(doc, '图2-3：考场管理页面', '显示考场列表、容量信息、操作按钮')
    
    # 2.4
    add_heading_custom(doc, '2.4 考生导入', level=2)
    doc.add_paragraph('通过 Excel 文件批量导入考生信息。')
    doc.add_paragraph('【操作步骤】')
    steps5 = [
        '点击【下载导入模板】获取标准格式',
        '按模板要求填写考生信息（姓名、学号、专业等）',
        '点击【选择文件】上传填写好的 Excel',
        '系统自动解析并显示导入预览',
        '确认无误后点击【确认导入】'
    ]
    for i, step in enumerate(steps5, 1):
        doc.add_paragraph(f'{i}. {step}')
    add_screenshot_placeholder(doc, '图2-4：考生导入页面', '显示模板下载按钮、文件上传区域、导入预览表格')
    
    doc.add_paragraph('【模板字段说明】')
    fields = [
        '姓名：考生真实姓名',
        '学号：研究生考生编号',
        '专业：报考专业名称',
        '学院：所属学院',
        '备注：特殊情况说明（可选）'
    ]
    for field in fields:
        doc.add_paragraph(field, style='List Bullet')
    
    # 2.5
    add_heading_custom(doc, '2.5 抽签执行', level=2)
    doc.add_paragraph('对指定批次的考生进行随机分组抽签。')
    doc.add_paragraph('【操作步骤】')
    steps6 = [
        '点击左侧菜单【抽签管理】',
        '选择要抽签的批次',
        '确认考生名单和考场信息',
        '点击【开始抽签】按钮',
        '系统随机分配考生到各考场组',
        '查看分组结果',
        '点击【导出结果】下载 Excel 文件'
    ]
    for i, step in enumerate(steps6, 1):
        doc.add_paragraph(f'{i}. {step}')
    add_screenshot_placeholder(doc, '图2-5：抽签执行页面', '显示考生列表、抽签按钮、分组结果预览')
    add_screenshot_placeholder(doc, '图2-6：抽签结果导出', '显示导出按钮、文件下载提示')
    
    doc.add_page_break()
    
    # ===== 第三章：志愿者操作指南 =====
    add_heading_custom(doc, '三、志愿者操作指南', level=1)
    
    add_heading_custom(doc, '3.1 选择考场', level=2)
    doc.add_paragraph('志愿者登录后，需先选择负责的考场。')
    doc.add_paragraph('【操作步骤】')
    steps7 = [
        '登录后进入考场选择页面',
        '查看分配的考场列表',
        '点击要负责的考场【进入】',
        '进入该考场的管理界面'
    ]
    for i, step in enumerate(steps7, 1):
        doc.add_paragraph(f'{i}. {step}')
    add_screenshot_placeholder(doc, '图3-1：志愿者考场选择页面', '显示可选择的考场卡片列表')
    
    add_heading_custom(doc, '3.2 查看分组信息', level=2)
    doc.add_paragraph('查看当前考场的考生分组和抽签顺序。')
    doc.add_paragraph('【功能说明】')
    features = [
        '查看本场考生名单',
        '查看分组情况和抽签顺序号',
        '确认考生到场状态',
        '记录实际抽签结果'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    add_screenshot_placeholder(doc, '图3-2：分组信息查看页面', '显示考生列表、组号、抽签顺序、状态标记')
    
    doc.add_page_break()
    
    # ===== 第四章：常见问题 =====
    add_heading_custom(doc, '四、常见问题', level=1)
    
    add_heading_custom(doc, '4.1 登录问题', level=2)
    
    doc.add_paragraph('Q1：统一身份认证跳转后显示"client_id为空"？')
    doc.add_paragraph('A：请清除浏览器缓存（Ctrl+Shift+R强制刷新）后重试，或联系管理员检查系统配置。', style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph('Q2：提示"账号不存在或已被禁用"？')
    doc.add_paragraph('A：请联系学院研究生教务老师确认账号是否已录入系统。', style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph('Q3：忘记系统账号密码？')
    doc.add_paragraph('A：请联系系统管理员重置密码。', style='List Bullet')
    
    add_heading_custom(doc, '4.2 操作问题', level=2)
    
    doc.add_paragraph('Q4：考生导入失败？')
    doc.add_paragraph('A：请检查：1）文件格式是否为.xlsx；2）必填字段是否完整；3）学号是否有重复。', style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph('Q5：抽签后能否重新抽？')
    doc.add_paragraph('A：已完成的抽签可以"重置"后重新执行，但会清除之前的结果，请谨慎操作。', style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph('Q6：如何修改已导入的考生信息？')
    doc.add_paragraph('A：在【考生管理】页面搜索该考生，点击【编辑】修改信息；或删除后重新导入。', style='List Bullet')
    
    # ===== 附录 =====
    doc.add_page_break()
    add_heading_custom(doc, '附录：技术支持', level=1)
    
    doc.add_paragraph('如遇到本手册无法解决的问题，请联系：')
    doc.add_paragraph()
    doc.add_paragraph('• 系统技术支持：XXX老师', style='List Bullet')
    doc.add_paragraph('• 联系电话：XXX-XXXXXXXX', style='List Bullet')
    doc.add_paragraph('• 办公地点：XXXX楼XXX室', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('• 学校信息中心（统一身份认证问题）：010-XXXXXXXX', style='List Bullet')
    
    # 保存文档
    output_path = '北科大研究生复试抽签系统-用户操作手册.docx'
    doc.save(output_path)
    print(f'手册已生成：{output_path}')

if __name__ == '__main__':
    main()
